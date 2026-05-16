from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "report.docx"


breed_classes = [
    "BUFFALO_JAFFARABADI",
    "BUFFALO_MEHSANA",
    "BUFFALO_MURRAH",
    "COW_GIR",
    "COW_HOLSTEIN",
    "COW_JERSEY",
    "COW_SAHIWAL",
]

breed_actual_cm = [
    [50, 1, 16, 1, 3, 0, 2],
    [9, 1, 5, 0, 0, 0, 0],
    [35, 2, 25, 1, 2, 0, 2],
    [9, 0, 2, 24, 1, 0, 39],
    [15, 0, 4, 0, 45, 1, 1],
    [22, 0, 9, 39, 17, 5, 53],
    [10, 0, 1, 19, 2, 2, 54],
]

breed_actual_metrics = [
    ["BUFFALO_JAFFARABADI", 50, 356, 100, 23, 0.3333, 0.6849, 0.4484],
    ["BUFFALO_MEHSANA", 1, 511, 3, 14, 0.2500, 0.0667, 0.1053],
    ["BUFFALO_MURRAH", 25, 425, 37, 42, 0.4032, 0.3731, 0.3876],
    ["COW_GIR", 24, 394, 60, 51, 0.2857, 0.3200, 0.3019],
    ["COW_HOLSTEIN", 45, 438, 25, 21, 0.6429, 0.6818, 0.6618],
    ["COW_JERSEY", 5, 381, 3, 140, 0.6250, 0.0345, 0.0654],
    ["COW_SAHIWAL", 54, 344, 97, 34, 0.3576, 0.6136, 0.4519],
]

disease_classes = ["FOOT_AND_MOUTH", "HEALTHY", "LUMPY"]
disease_actual_cm = [
    [0, 0, 0],
    [0, 0, 0],
    [5, 40, 603],
]
disease_actual_metrics = [
    ["FOOT_AND_MOUTH", 0, 643, 5, 0, 0.0000, 0.0000, 0.0000],
    ["HEALTHY", 0, 608, 40, 0, 0.0000, 0.0000, 0.0000],
    ["LUMPY", 603, 0, 0, 45, 1.0000, 0.9306, 0.9640],
]

breed_target_metrics = [
    ["BUFFALO_JAFFARABADI", 66, 449, 7, 7, 0.9041, 0.9041, 0.9041],
    ["BUFFALO_MEHSANA", 14, 513, 1, 1, 0.9333, 0.9333, 0.9333],
    ["BUFFALO_MURRAH", 60, 456, 6, 7, 0.9091, 0.8955, 0.9023],
    ["COW_GIR", 68, 446, 8, 7, 0.8947, 0.9067, 0.9007],
    ["COW_HOLSTEIN", 60, 458, 5, 6, 0.9231, 0.9091, 0.9160],
    ["COW_JERSEY", 131, 369, 15, 14, 0.8973, 0.9034, 0.9003],
    ["COW_SAHIWAL", 79, 432, 9, 9, 0.8977, 0.8977, 0.8977],
]

disease_target_metrics = [
    ["FOOT_AND_MOUTH", 199, 417, 15, 17, 0.9299, 0.9213, 0.9256],
    ["HEALTHY", 202, 418, 14, 14, 0.9352, 0.9352, 0.9352],
    ["LUMPY", 208, 420, 12, 8, 0.9455, 0.9630, 0.9541],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill="1F4E79", zebra=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "80")
                node.set(qn("w:type"), "dxa")
            if r == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            elif zebra and r % 2 == 0:
                set_cell_shading(cell, "F3F6F8")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    return p


def add_metric_table(doc, title, rows):
    add_heading(doc, title, 3)
    headers = ["Class", "TP", "TN", "FP", "FN", "Precision", "Recall", "F1 Score"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                value = f"{value:.4f}"
            set_cell_text(cells[i], value, size=7.8, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    style_table(table)
    doc.add_paragraph()


def add_confusion_table(doc, title, labels, cm):
    add_heading(doc, title, 3)
    short = {
        "BUFFALO_JAFFARABADI": "Jaffarabadi",
        "BUFFALO_MEHSANA": "Mehsana",
        "BUFFALO_MURRAH": "Murrah",
        "COW_GIR": "Gir",
        "COW_HOLSTEIN": "Holstein",
        "COW_JERSEY": "Jersey",
        "COW_SAHIWAL": "Sahiwal",
        "FOOT_AND_MOUTH": "F&M",
        "HEALTHY": "Healthy",
        "LUMPY": "Lumpy",
    }
    table = doc.add_table(rows=1, cols=len(labels) + 1)
    header = table.rows[0].cells
    set_cell_text(header[0], "Actual / Predicted", bold=True, size=7.5)
    for i, label in enumerate(labels, start=1):
        set_cell_text(header[i], short.get(label, label), bold=True, size=7.5)
    for label, values in zip(labels, cm):
        cells = table.add_row().cells
        set_cell_text(cells[0], short.get(label, label), size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, value in enumerate(values, start=1):
            set_cell_text(cells[i], value, size=7.8)
    style_table(table)
    doc.add_paragraph()


def add_summary_table(doc, title, rows):
    add_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=5)
    headers = ["Model", "Accuracy", "Precision", "Recall", "F1 Score"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    style_table(table)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PashuDrishti AI Model Evaluation Report")
run.font.name = "Arial"
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Breed and Disease Classification Models")
run.font.name = "Arial"
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(80, 80, 80)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Includes data preprocessing, feature extraction, confusion matrices, and class-wise/model-wise metrics.")
run.font.name = "Arial"
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()
add_heading(doc, "Data Preprocessing", 1)
for text in [
    "The raw cow and buffalo images are arranged in class-wise folders. TensorFlow reads these folders and automatically assigns numeric labels based on the folder names.",
    "All input images are resized to 224 x 224 pixels so that every sample has the same shape before entering the CNN model.",
    "Images are processed in batches, commonly with a batch size of 32, so the model can train efficiently on groups of images.",
    "For the disease model, the dataset is split into training and validation data using an 80:20 split. This helps evaluate model behavior on unseen validation images.",
    "Pixel values are normalized from the original 0-255 range to a smaller 0-1 range where applicable. Normalization helps stabilize and speed up model learning.",
    "Data augmentation techniques such as rotation, zoom, horizontal flipping, shifting, and brightness variation can be applied during training to improve generalization.",
]:
    add_body(doc, text)

add_heading(doc, "Feature Extraction", 1)
for text in [
    "Feature extraction is performed by a Convolutional Neural Network. The CNN learns useful visual patterns directly from the input images.",
    "Early convolution layers identify simple features such as edges, corners, colors, textures, and local shapes.",
    "Deeper layers learn higher-level features such as breed-specific body structure, horn shape, coat pattern, skin texture, lumps, mouth or foot infection marks, and healthy skin appearance.",
    "Pooling layers reduce spatial size while preserving important patterns, which helps make the feature representation more compact.",
    "The extracted features are converted into a feature vector and passed to the final classification layer. The softmax output gives the predicted breed or disease class.",
]:
    add_body(doc, text)

add_heading(doc, "Actual Evaluation Results", 1)
add_body(doc, "The following values are taken from the saved metric files in the project. The disease model accuracy appears high, but the validation set used in the saved result contains only true LUMPY samples, so macro scores are more informative than accuracy.")

add_summary_table(
    doc,
    "Model-wise Actual Metrics",
    [
        ["Breed Model", "0.3856", "0.4140 macro / 0.4557 weighted", "0.3964 macro / 0.3856 weighted", "0.3460 macro / 0.3324 weighted"],
        ["Disease Model", "0.9306", "0.3333 macro / 1.0000 weighted", "0.3102 macro / 0.9306 weighted", "0.3213 macro / 0.9640 weighted"],
    ],
)

add_metric_table(doc, "Breed Model Class-wise Actual Metrics", breed_actual_metrics)
add_confusion_table(doc, "Breed Model Actual Confusion Matrix", breed_classes, breed_actual_cm)

add_metric_table(doc, "Disease Model Class-wise Actual Metrics", disease_actual_metrics)
add_confusion_table(doc, "Disease Model Actual Confusion Matrix", disease_classes, disease_actual_cm)

doc.add_page_break()
add_heading(doc, "Alternative Good-Model Target Values", 1)
add_body(doc, "The values below are alternative good-model target values for reporting an expected or improved model. They should not be presented as measured results unless the model is retrained and evaluated to produce them.")

add_summary_table(
    doc,
    "Model-wise Target Metrics",
    [
        ["Breed Model", "0.9036", "Around 0.91", "Around 0.91", "Around 0.91"],
        ["Disease Model", "0.9398", "Around 0.94", "Around 0.94", "Around 0.94"],
    ],
)
add_metric_table(doc, "Breed Model Class-wise Target Metrics", breed_target_metrics)
add_metric_table(doc, "Disease Model Class-wise Target Metrics", disease_target_metrics)

add_heading(doc, "Conclusion", 1)
add_body(doc, "The breed model requires improvement because its actual accuracy and macro F1 score are low. The disease model requires a better balanced validation dataset because the saved evaluation result is dominated by the LUMPY class. A reliable final evaluation should use a balanced test set with samples from every class.")

doc.save(OUT)
print(OUT)
